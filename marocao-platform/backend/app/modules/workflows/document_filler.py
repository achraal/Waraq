import logging
from pathlib import Path
from typing import Dict, Any
from docx import Document
from openpyxl import load_workbook

logger = logging.getLogger(__name__)

class DocumentFiller:
    """Moteur d'injection des données dans les documents."""
    def fill_docx(self,input_path: str | Path,output_path: str | Path,values: Dict[str, Any]) -> Path:
        """
        Remplit les placeholders présents dans un fichier DOCX.
        """

        input_path = Path(input_path)
        output_path = Path(output_path)
        document = Document(str(input_path))
        self._replace_docx_paragraphs(document,values)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_paragraph(paragraph,values)
        document.save(str(output_path))
        logger.info("DOCX rempli | output=%s",output_path)
        return output_path

    def _replace_docx_paragraphs(self,document: Document,values: Dict[str, Any]):
        """Remplace les placeholders des paragraphes DOCX."""
        for paragraph in document.paragraphs:
            self._replace_paragraph(paragraph,values)

    def _replace_paragraph(self,paragraph,values: Dict[str, Any]):
        """
        Remplace les placeholders d'un paragraphe.
        Cette implémentation traite le texte consolidé du paragraphe.
        """

        text = paragraph.text
        for key, value in values.items():
            placeholder = f"{{{{key}}}}"
            text = text.replace(placeholder,"" if value is None else str(value))
        if text != paragraph.text:
            paragraph.text = text

    def fill_xlsx(self, input_path: str | Path,output_path: str | Path,values: Dict[str, Any]) -> Path:
        """
        Remplit les cellules Excel contenant des placeholders.
        """

        workbook = load_workbook(filename=str(input_path))
        for worksheet in workbook.worksheets:
            for row in worksheet.iter_rows():
                for cell in row:
                    if not isinstance(cell.value,str):
                        continue
                    value = cell.value
                    for key, replacement in values.items():
                        placeholder = f"{{{{key}}}}"
                        value = value.replace(placeholder,"" if replacement is None else str(replacement))
                    cell.value = value
        workbook.save(str(output_path))
        logger.info("XLSX rempli | output=%s",output_path)
        return output_path