import logging
from pathlib import Path
from typing import Dict
from uuid import UUID
from docx import Document
from sqlalchemy.orm import Session
from backend.app.database.models import User, Tender, CompanyProfile

logger = logging.getLogger(__name__)

class AdminDocGeneratorService:
    """Générateur des documents administratifs Waraq."""
    def __init__(self, templates_dir: Path):
        self.templates_dir = Path(templates_dir)

    def _get_context(self, profile: CompanyProfile, tender: Tender) -> Dict[str, str]:
        """
        Construit le dictionnaire des variables métier utilisées dans les documents administratifs.
        """
        return {
            "RAISON_SOCIALE": getattr(profile,"company_name",None) or profile.manager_name,
            "NOM_GERANT": profile.manager_name,
            "ADRESSE": profile.address,
            "ICE": profile.ice or "",
            "RC_NUM": getattr(profile,"rc_number",None) or getattr(profile,"cin_number",""),
            "RIB": profile.rib,
            "BANQUE": profile.bank_name,
            "OBJET_AO": tender.title,
            "REF_AO": tender.reference,
            "BUYER": tender.buyer,
        }

    def _replace_document(self, document: Document, values: Dict[str, str]):
        """
        Remplace les variables dans les paragraphes et tableaux DOCX.
        """
        for paragraph in document.paragraphs:
            self._replace_paragraph(paragraph, values)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_paragraph(paragraph, values)

    def _replace_paragraph(self, paragraph, values: Dict[str, str]):
        """Remplace les placeholders d'un paragraphe."""
        text = paragraph.text
        for key, value in values.items():
            text = text.replace("{{" + key + "}}",str(value or ""))
        if text != paragraph.text:
            paragraph.text = text

    def _generate(self,db: Session,user_id: UUID,tender_id: UUID,template_name: str,output_path: Path) -> Path:
        """
        Génère un document administratif à partir d'un modèle Waraq.
        """
        user = (db.query(User).filter(User.id == user_id).first())
        tender = (db.query(Tender).filter(Tender.id == tender_id).first())

        if not user:
            raise ValueError("Utilisateur introuvable.")
        if not user.company_profile:
            raise ValueError("Profil entreprise introuvable.")
        if not tender:
            raise ValueError("Appel d'offres introuvable.")
        template_path = (self.templates_dir /template_name)
        if not template_path.exists():
            raise FileNotFoundError(f"Modèle introuvable : {template_path}")
        output_path.parent.mkdir(parents=True,exist_ok=True)
        document = Document(str(template_path))
        values = self._get_context(user.company_profile,tender)
        self._replace_document(document,values)
        document.save(str(output_path))

        logger.info(
            "Document administratif généré | "
            "type=%s | tender=%s | output=%s",
            template_name,
            tender.reference,
            output_path
        )
        return output_path

    def generate_acte_engagement(self,db: Session,user_id: UUID,tender_id: UUID,output_path: Path) -> Path:
        """
        Génère l'Acte d'Engagement Waraq.
        """
        return self._generate(
            db=db,
            user_id=user_id,
            tender_id=tender_id,
            template_name="acte_engagement_waraq_template.docx",
            output_path=output_path
        )

    def generate_declaration_honneur(self,db: Session,user_id: UUID,tender_id: UUID,output_path: Path) -> Path:
        """
        Génère la Déclaration sur l'Honneur Waraq.
        """
        return self._generate(
            db=db,
            user_id=user_id,
            tender_id=tender_id,
            template_name="declaration_honneur_waraq_template.docx",
            output_path=output_path
        )