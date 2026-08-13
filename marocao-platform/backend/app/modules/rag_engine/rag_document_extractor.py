import os, re, shutil, logging, subprocess, tempfile, cv2, fitz, unicodedata
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
import numpy as np
import pandas as pd
from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak)
from rapidocr_onnxruntime import RapidOCR
from backend.app.config import settings

logger = logging.getLogger("waraq.rag.document_extractor")

BASE_STORAGE_DIR = Path(r"C:\Users\achra\Desktop\Intern\Project\marocao-platform\data_storage")

ZONE_SIGNATURE_KEYWORDS = ["signature","signatures","cachet","cachets","signé","signe","signée","signe par","signé par","signature et cachet","cachet et signature"]

ZONE_VALIDATION_KEYWORDS = ["lu et accepte","lu et accepté","lu et approuve","lu et approuvé","vu et présente","vu et présenté","vu et verifie","vu et vérifie","vise par","visé par","visa","approuve par","approuvé par","approuvé","visa du","visa de"]

ZONE_DATE_KEYWORDS = ["fait a","fait à","date"]

ZONE_ADMINISTRATIVE_KEYWORDS = ZONE_SIGNATURE_KEYWORDS + ZONE_VALIDATION_KEYWORDS

# OCR RAPIDOCR

_rapid_ocr_instance = None
def _get_onnx_ocr_engine():
    global _rapid_ocr_instance
    if _rapid_ocr_instance is None:
        try:
            _rapid_ocr_instance = RapidOCR()
            logger.info("[RAG-OCR] RapidOCR / ONNX Runtime initialisé.")
        except Exception as e:
            logger.error("[RAG-OCR] Impossible d'initialiser RapidOCR : %s", e, exc_info=True)
            return None
    return _rapid_ocr_instance

def _ocr_image(engine, image) -> str:
    try:
        result, _ = engine(image)
        lignes = []
        if result:
            for item in result:
                texte = str(item[1]).strip()
                confiance = float(item[2])
                if texte and confiance >= 0.40:
                    lignes.append(texte)
        return "\n".join(lignes).strip()
    except Exception as e:
        logger.error("[RAG-OCR] Erreur OCR image : %s", e)
        return ""

def _ocr_est_de_mauvaise_qualite(texte: str) -> bool:
    if not texte:
        return True
    texte = texte.strip()
    if len(texte) < 250:
        return True
    nb_lettres = len(re.findall(r"[A-Za-zÀ-ÿ\u0600-\u06FF]", texte))

    ratio = nb_lettres / max(len(texte), 1)
    if ratio < 0.60:
        return True
    parasites = len(re.findall(r"[+*#=<>\\|~]", texte))

    if parasites >= 2:
        return True
    return False

def extraire_page_complete_fast_ocr(page, page_number: int) -> str:
    """
    OCR COMPLET d'une page. Contrairement à l'ancien OCR d'en-tête,cette fonction analyse toute la page.
    """
    try:
        mat = fitz.Matrix(2.0, 2.0)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img = np.frombuffer(pix.samples, dtype=np.uint8).reshape((pix.height, pix.width, 3))
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        engine = _get_onnx_ocr_engine()
        if engine is None:
            return ""
        texte = _ocr_image(engine, gray)
        logger.info("[RAG-OCR] Page %s : %s caractères extraits.", page_number, len(texte))
        return texte
    except Exception as e:
        logger.error("[RAG-OCR] Erreur page %s : %s", page_number, e, exc_info=True)
        return ""

DOCUMENT_PATTERNS = {
    "CURRICULUM_VITAE": [
        r"\bcurriculum\s+vitae\b",
        r"\bexpérience\s+professionnelle\b",
        r"\bparcours\s+professionnel\b",
    ],

    "DECLARATION_HONNEUR": [
        r"d[ée]claration\s+sur\s+l['’]honneur",
        r"d[ée]claration\s+sur\s+honneur",
    ],

    "ACTE_ENGAGEMENT": [
        r"\bacte\s+d['’]engagement\b",
        r"\bacte\s+d\s+engagement\b",
    ],

    "DECLARATION_IDENTITE": [
        r"d[ée]claration\s+d['’]identit[ée]",
        r"identification\s+du\s+soumissionnaire",
        r"identification\s+du\s+candidat",
    ],

    "BDP": [
        r"\bbordereau\s+des\s+prix\b",
        r"\bbordereau\s+de\s+prix\b",
        r"d[ée]tail\s+estimatif",
        r"d[ée]tail\s+des\s+prix",
    ],

    "BPU": [
        r"\bbordereau\s+des\s+prix\s+unitaires\b",
        r"\bbpu\b",
    ],

    "DPGF": [
        r"d[ée]composition\s+du\s+prix\s+global",
        r"d[ée]composition\s+des\s+prix",
        r"\bdpgf\b",
    ],

    "RC": [
        r"\br[èe]glement\s+de\s+la?\s+consultation\b",
        r"\br[èe]glement\s+de\s+consultation\b",
    ],

    "CPS": [
        r"\bcahier\s+des\s+prescriptions\s+sp[ée]ciales\b",
    ],

    "DECLARATION_FISCALE": [
        r"\battestation\s+fiscale\b",
        r"\bsituation\s+fiscale\b",
    ],

    "DECLARATION_SOCIALE": [
        r"\battestation\s+cnss\b",
        r"\bcaisse\s+nationale\s+de\s+s[ée]curit[ée]\s+sociale\b",
    ],
}

def normaliser_texte_modele(text: str) -> str:
    """
    Normalisation légère utilisée uniquement pour détecter les signatures
    fortes des modèles.
    """
    if not text:
        return ""

    text = unicodedata.normalize("NFD", text)
    text = "".join(c for c in text if unicodedata.category(c) != "Mn"    )
    text = text.upper()
    text = text.replace("’", "'")
    text = text.replace("–", "-")
    text = text.replace("—", "-")
    text = re.sub(r"\s+", " ", text)
    return text.strip()

SIGNATURES_FORTES_MODELES = {
    "BDP": [
        "BORDEREAU DES PRIX",
        "BORDEREAU DES PRIX DETAIL ESTIMATIF",
        "BORDEREAU DES PRIX DETAIL ESTIMATIF",
        "BORDEREAU DE PRIX",
        "BORDEREAU DES PRIX UNITAIRES",
        "PRIX UNITAIRES",
        "DESIGNATION",
        "UNITE",
        "QUANTITE",
        "PRIX UNITAIRE",
    ],
    "ACTE_ENGAGEMENT": [
        "ACTE D'ENGAGEMENT",
        "ACTE D ENGAGEMENT",
        "ACTE D’ENGAGEMENT",
        "JE SOUSSIGNE",
        "JE SOUSSIGNE",
        "AGISSANT AU NOM ET POUR LE COMPTE",
        "M'ENGAGE",
        "M ENGAGE",
        "S'ENGAGE",
        "S ENGAGE",
    ],
    "DECLARATION_HONNEUR": [
        "DECLARATION SUR L'HONNEUR",
        "DECLARATION SUR L HONNEUR",
        "DECLARATION D'HONNEUR",
        "DECLARATION SUR LHONNEUR",
    ],
}

def page_est_vrai_modele(texte_page: str, model_type: str) -> bool:
    """
    Détermine si une page contient réellement le modèle demandé.
    IMPORTANT : Une simple mention du modèle dans une clause n'est pas suffisante.
    """

    texte = normaliser_texte_modele(texte_page)
    if not texte:
        return False
    signatures = SIGNATURES_FORTES_MODELES.get(model_type, [])

    if not signatures:
        return False
    score = 0
    signatures_trouvees = []

    for signature in signatures:
        signature_norm = normaliser_texte_modele(signature)
        if signature_norm in texte:
            score += 1
            signatures_trouvees.append(signature)
            
    if model_type == "BDP":

        # =========================================================
        # BDP = véritable tableau de prix
        #
        # Une mention de :
        #   "bordereau des prix"
        #   "prix"
        #   "montant"
        #
        # ne suffit jamais.
        # =========================================================
        # 1. Structure narrative forte

        marqueurs_narratifs = [
            "ARTICLE ",
            "CHAPITRE ",
            "DOCUMENTS CONSTITUTIFS",
            "REFERENCE AUX TEXTES",
            "REFERENCES AUX TEXTES",
            "PENALITES",
            "MODALITES ET CONDITIONS DE LIVRAISON",
            "CONDITIONS DE LIVRAISON",
            "DELAI D'EXECUTION",
            "DELAI D EXECUTION",
            "GARANTIE",
            "OBLIGATIONS DU FOURNISSEUR",
            "OBLIGATIONS DU TITULAIRE",
            "CONDITIONS DE PAIEMENT",
            "REGLEMENT DES FACTURES",
        ]

        nombre_narratif = sum(1 for marqueur in marqueurs_narratifs if marqueur in texte)
        # 2. Colonnes spécifiques à un tableau de prix
        colonnes_fortes = [
            "DESIGNATION",
            "DESIGNATION DES PRESTATIONS",
            "UNITE",
            "QUANTITE",
            "PRIX UNITAIRE",
            "PRIX UNITAIRE HT",
            "PRIX UNITAIRE TTC",
            "MONTANT",
            "MONTANT HT",
            "MONTANT TTC",
            "TOTAL HT",
            "TOTAL TTC",
        ]
        colonnes_presentes = []

        for colonne in colonnes_fortes:
            if normaliser_texte_modele(colonne) in texte:
                colonnes_presentes.append(normaliser_texte_modele(colonne))
        colonnes_uniques = set(colonnes_presentes)

        # 3. Détection de lignes tabulaires

        lignes = [ligne.strip() for ligne in texte_page.splitlines() if ligne.strip()]
        lignes_avec_nombres = 0
        for ligne in lignes:
            nombres_ligne = re.findall(r"\b\d+(?:[.,]\d+)?\b", ligne)
            if len(nombres_ligne) >= 2:
                lignes_avec_nombres += 1

        # 4. Séparateurs de colonnes
        separateurs_tableau = texte_page.count("|") + texte_page.count("\t")

        # 5. Détection d'un titre BDP
        titre_bdp = any(titre in texte for titre in [
                "BORDEREAU DES PRIX",
                "BORDEREAU DE PRIX",
                "DETAIL ESTIMATIF",
                "DETAIL DES PRIX",
                "BORDEREAU DES PRIX DETAIL ESTIMATIF",
            ]
        )

        # REJET ABSOLU DES PAGES NARRATIVES
        if nombre_narratif >= 1:
            # Une page contenant un ARTICLE + des mots financiers reste une page de CPS/RC et non un BDP.
            if not titre_bdp or len(colonnes_uniques) < 3:
                logger.info(
                    "[RAG-MODELES][BDP] REJET page narrative | " "narratif=%d | colonnes=%s | titre=%s",
                    nombre_narratif,list(colonnes_uniques),titre_bdp)

                return False

        # =========================================================
        # VALIDATION 1
        # Titre BDP + au moins 3 colonnes fortes
        # =========================================================

        if titre_bdp and len(colonnes_uniques) >= 3:
            logger.info("[RAG-MODELES][BDP] CONFIRME | " "titre + 3 colonnes | %s",list(colonnes_uniques))
            return True

        # =========================================================
        # VALIDATION 2
        #
        # 3 colonnes fortes + plusieurs lignes numériques
        #
        # Même sans titre explicite, cela peut être une continuation
        # ou un tableau BDP.
        # =========================================================

        if len(colonnes_uniques) >= 3 and lignes_avec_nombres >= 2:
            logger.info(
                "[RAG-MODELES][BDP] CONFIRME | "
                "structure tabulaire | colonnes=%s | "
                "lignes_numeriques=%d",
                list(colonnes_uniques),
                lignes_avec_nombres
            )
            return True

        # =========================================================
        # VALIDATION 3
        #
        # 2 colonnes fortes seulement si vraie structure tabulaire
        # =========================================================

        if len(colonnes_uniques) >= 2 and lignes_avec_nombres >= 3 and separateurs_tableau >= 2:
            logger.info("[RAG-MODELES][BDP] CONFIRME | " "structure tabulaire forte.")
            return True

        # Titre seul = JAMAIS suffisant

        logger.info(
            "[RAG-MODELES][BDP] REJET | " "titre=%s | colonnes=%s | " "lignes_numeriques=%d | separateurs=%d",
            titre_bdp,list(colonnes_uniques),lignes_avec_nombres,separateurs_tableau)
        return False

    # ACTE D'ENGAGEMENT
    if model_type == "ACTE_ENGAGEMENT": 
        # Il faut une signature explicite.
        if ("ACTE D'ENGAGEMENT" in texte or "ACTE D ENGAGEMENT" in texte):
            return True

        # Ou une combinaison très forte de formulations caractéristiques de l'acte.
        indices_forts = 0

        for motif in [
            "JE SOUSSIGNE",
            "AGISSANT AU NOM ET POUR LE COMPTE",
            "M'ENGAGE",
            "M ENGAGE",
        ]:
            if motif in texte:
                indices_forts += 1
        return indices_forts >= 2

    # DECLARATION SUR L'HONNEUR
    if model_type == "DECLARATION_HONNEUR":
        titres = [
            "DECLARATION SUR L'HONNEUR",
            "DECLARATION D'HONNEUR",
            "DECLARATION SUR LHONNEUR",
        ]

        titre_trouve = any(titre in texte for titre in titres)
        if not titre_trouve:
            return False

        # Le titre doit être accompagné d'au moins un élément caractéristique d'une véritable déclaration.

        indices_declaration = [
            "JE SOUSSIGNE",
            "JE SOUSSIGNEE",
            "DECLARE SUR L'HONNEUR",
            "ATTESTE SUR L'HONNEUR",
            "CERTIFIE SUR L'HONNEUR",
            "FAIT A",
            "FAIT À",
            "SIGNATURE",
        ]

        indices_trouves = sum(1 for indice in indices_declaration if normaliser_texte_modele(indice) in texte)
        # Titre + au moins 1 indice administratif.
        return indices_trouves >= 1

def trouver_page_debut_modele(pages: List[str], model_type: str,) -> Optional[int]:

    for index, texte_page in enumerate(pages):
        if page_est_vrai_modele(texte_page,model_type):
            logger.info("[RAG-MODELES] Début confirmé %s : page %d",model_type,index + 1)
            return index

    logger.warning("[RAG-MODELES] Aucun début confirmé pour %s", model_type)
    return None

def trouver_page_fin_modele(pages: List[str], model_type: str,start_index: int,) -> int:

    if start_index >= len(pages):
        return start_index

    # BDP
    if model_type == "BDP":
        last_valid_page = start_index
        for i in range(start_index + 1, min(start_index + 30, len(pages))):
            texte = normaliser_texte_modele(pages[i])
            if not texte:
                last_valid_page = i
                continue
            # 1. Détection d'un nouveau modèle fort
            nouveau_modele = False
            signatures_nouveaux_documents = [
                "ACTE D'ENGAGEMENT",
                "DECLARATION SUR L'HONNEUR",
                "DECLARATION D'HONNEUR",
                "REGLEMENT DE LA CONSULTATION",
                "CAHIER DES PRESCRIPTIONS SPECIALES",
                "CURRICULUM VITAE",
                "DECLARATION D'IDENTITE",
            ]

            for signature in signatures_nouveaux_documents:
                if signature in texte:
                    nouveau_modele = True
                    logger.info("[RAG-MODELES] Fin BDP détectée avant page %d : %s",i + 1,signature)
                    break
            if nouveau_modele:
                break
            # 2. Vérification structure BDP

            if page_est_continuation_bdp(texte):
                last_valid_page = i
                continue
            # 3. Page ambiguë: Une page juste après le BDP peut être pauvre. On la conserve seulement si elle ressemble encore à une page de tableau.
            lignes = [ligne.strip() for ligne in texte.splitlines() if ligne.strip()]
            nombres = re.findall(r"\b\d+(?:[.,]\d+)?\b", texte)

            if len(lignes) >= 3 and len(nombres) >= 2:
                last_valid_page = i
                continue
            # Sinon on considère que le BDP est terminé.
            break
        return last_valid_page
    
    # ACTE ENGAGEMENT
    if model_type == "ACTE_ENGAGEMENT":
        last_valid_page = start_index
        for i in range(start_index, min(start_index + 15, len(pages))):
            texte = normaliser_texte_modele(pages[i])

            # Une page appartient à l'acte si elle contient des marqueurs caractéristiques.
            marqueurs = [
                "ACTE D'ENGAGEMENT",
                "ACTE D ENGAGEMENT",
                "JE SOUSSIGNE",
                "AGISSANT AU NOM ET POUR LE COMPTE",
                "M'ENGAGE",
                "M ENGAGE",
                "LU ET ACCEPTE",
                "LE FOURNISSEUR",
                "LE TITULAIRE",
                "SIGNATURE",
            ]

            score = sum(1 for m in marqueurs if m in texte)

            # Une page suffisamment caractéristique reste dans le modèle.
            if score >= 1:
                last_valid_page = i
            else:
                # On tolère UNE page intermédiaire pauvre.
                if i == last_valid_page + 1:
                    continue
                break
        return last_valid_page

    # DECLARATION HONNEUR
    if model_type == "DECLARATION_HONNEUR":
        last_valid_page = start_index
        for i in range(start_index, min(start_index + 5, len(pages))):
            texte = normaliser_texte_modele(pages[i])
            marqueurs = [
                "DECLARATION SUR L'HONNEUR",
                "DECLARATION SUR L HONNEUR",
                "JE SOUSSIGNE",
                "DECLARE SUR L'HONNEUR",
                "ATTESTE SUR L'HONNEUR",
            ]
            score = sum(1 for m in marqueurs if m in texte)
            if score >= 1:
                last_valid_page = i
            else:
                break
        return last_valid_page
    return start_index

def determiner_plage_modele(pages: List[str],model_type: str,) -> Optional[Tuple[int, int]]:

    start_index = trouver_page_debut_modele(pages=pages,model_type=model_type,)
    if start_index is None:
        logger.warning("[RAG-MODELES] Modèle %s ignoré : " "aucune page de début confirmée.", model_type)
        return None
    end_index = trouver_page_fin_modele(pages=pages, model_type=model_type,start_index=start_index,)
    logger.info("[RAG-MODELES] Plage finale %s : pages %d -> %d",model_type,start_index + 1,end_index + 1)
    return start_index, end_index

def page_est_continuation_bdp(texte_page: str) -> bool:
    texte = normaliser_texte_modele(texte_page)

    # Une continuation BDP ne doit pas être une page narrative

    marqueurs_narratifs = [
        "ARTICLE ",
        "CHAPITRE ",
        "PENALITES",
        "MODALITES ET CONDITIONS",
        "CONDITIONS DE LIVRAISON",
        "GARANTIE",
        "OBLIGATIONS DU FOURNISSEUR",
        "OBLIGATIONS DU TITULAIRE",
    ]

    if any(marqueur in texte for marqueur in marqueurs_narratifs):
        return False

    colonnes = [
        "DESIGNATION",
        "DESIGNATION DES PRESTATIONS",
        "UNITE",
        "QUANTITE",
        "PRIX UNITAIRE",
        "PRIX UNITAIRES",
        "PRIX UNITAIRE HT",
        "PRIX UNITAIRE TTC",
        "MONTANT",
        "MONTANT HT",
        "MONTANT TTC",
    ]

    colonnes_presentes = sum(1 for colonne in colonnes if normaliser_texte_modele(colonne) in texte)
    lignes = [ligne.strip() for ligne in texte_page.splitlines() if ligne.strip()]
    lignes_numeriques = sum(1 for ligne in lignes if len(re.findall(r"\b\d+(?:[.,]\d+)?\b", ligne)) >= 2)
    # Une continuation doit présenter une structure de tableau.
    return colonnes_presentes >= 2 and lignes_numeriques >= 2

TYPES_MODELES = {"CURRICULUM_VITAE","DECLARATION_HONNEUR","ACTE_ENGAGEMENT","DECLARATION_IDENTITE","BDP","BPU","DPGF"}

def normaliser_texte_detection(text: str) -> str:
    text = text.lower()
    text = text.replace("’", "'").replace("`", "'")
    text = re.sub(r"\s+", " ", text)
    return text

def detecter_types_documents(text: str) -> List[str]:
    text_norm = normaliser_texte_detection(text)
    types_detectes = []
    for document_type, patterns in DOCUMENT_PATTERNS.items():
        for pattern in patterns:
            if re.search(pattern, text_norm, re.IGNORECASE):
                types_detectes.append(document_type)
                break
    return types_detectes

def _extraire_pdf(file_path: str) -> Dict[str, Any]:
    pages = []
    doc = fitz.open(file_path)
    try:
        for page_index in range(len(doc)):
            page = doc.load_page(page_index)
            page_number = page_index + 1

            # 1. EXTRACTION TEXTE NATIVE
            texte = page.get_text("text",sort=True).strip()
            extraction_method = "PYMUPDF"

            # 2. SI PAGE SCANNEE -> OCR COMPLET
            if not texte or _ocr_est_de_mauvaise_qualite(texte):
                logger.info("[RAG] Page %s : texte natif absent/faible -> OCR COMPLET.", page_number)
                texte_ocr = extraire_page_complete_fast_ocr(page, page_number)
                if len(texte_ocr) > len(texte):
                    texte = texte_ocr
                    extraction_method = "RAPIDOCR_ONNX_FULL_PAGE"
            pages.append({"page_number": page_number,"text": texte,"method": extraction_method})
    finally:
        doc.close()
    texte_complet = "\n\n".join( (f"===== PAGE {p['page_number']} =====\n"f"{p['text']}") for p in pages if p["text"].strip())
    return {"text": texte_complet,"pages": pages,"page_count": len(pages),}

def _extraire_docx(file_path: str) -> Dict[str, Any]:
    document = Document(file_path)
    lignes = []
    # Paragraphes
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lignes.append(text)
    # Tableaux
    for table_index, table in enumerate(document.tables, start=1):
        lignes.append(f"\n===== TABLEAU {table_index} =====")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lignes.append(" | ".join(cells))
    texte = "\n".join(lignes)
    return {"text": texte,"pages": [],"page_count": None,}

def _convertir_avec_libreoffice(file_path: str,output_dir: Optional[str] = None,) -> Optional[str]:
    """
    Convertit DOC/DOCX en PDF avec LibreOffice. Si output_dir est fourni : le PDF est conservé dans ce dossier.
    Sinon : un dossier temporaire est utilisé.
    """
    temp_dir = None
    try:
        if output_dir:
            output_path = Path(output_dir)
            output_path.mkdir(parents=True, exist_ok=True)
            conversion_dir = output_path
        else:
            temp_dir = tempfile.mkdtemp(prefix="waraq_rag_convert_")
            conversion_dir = Path(temp_dir)
        conversion_dir.mkdir(parents=True, exist_ok=True)
        logger.info("[RAG-CONVERT] Conversion DOC/DOCX -> PDF : %s",file_path)

        result = subprocess.run(
            [
                settings.LIBREOFFICE_PATH,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(conversion_dir),
                str(file_path),
            ],
            stdout=subprocess.PIPE,stderr=subprocess.PIPE,text=True,timeout=180,)

        logger.info("[RAG-CONVERT] stdout=%s",result.stdout.strip())
        if result.stderr:
            logger.warning("[RAG-CONVERT] stderr=%s",result.stderr.strip())
        if result.returncode != 0:
            logger.error("[RAG-CONVERT] Échec conversion | returncode=%s | stderr=%s",result.returncode,result.stderr,)
            return None
        expected_pdf = (conversion_dir / f"{Path(file_path).stem}.pdf")
        if not expected_pdf.exists():
            logger.error("[RAG-CONVERT] PDF attendu introuvable : %s",expected_pdf)
            return None
        logger.info("[RAG-CONVERT] PDF généré avec succès : %s",expected_pdf)
        return str(expected_pdf)
    except subprocess.TimeoutExpired:
        logger.error("[RAG-CONVERT] Timeout LibreOffice : %s",file_path)
        return None

    except Exception as e:
        logger.error("[RAG-CONVERT] Erreur conversion %s : %s",file_path,e,exc_info=True,)
        return None

    finally:
        # On ne supprime PAS output_dir.
        # Le dossier temporaire est volontairement conservé ici
        # seulement pour éviter de casser le flux existant.
        #
        # La version utilisée par le RAG fournira toujours output_dir.
        pass

def _extraire_excel(file_path: str) -> Dict[str, Any]:
    extension = ( Path(file_path).suffix.lower() )
    engine = None
    if extension == ".xls":
        engine = "xlrd"
    elif extension in [".xlsx",".xlsm",]:
        engine = "openpyxl"
    sheets = pd.read_excel(file_path, sheet_name=None, header=None, engine=engine, dtype=object,)
    lignes = []
    feuilles = []
    for sheet_name, dataframe in sheets.items():
        lignes.append(f"\n===== FEUILLE : {sheet_name} =====")
        rows = []
        for _, row in dataframe.iterrows():
            values = []
            for value in row.tolist():
                if pd.isna(value):
                    value = ""
                else:
                    value = str(value).strip()
                values.append(value)

            # Evite les lignes totalement vides
            if any(values):
                ligne = " | ".join(values)
                lignes.append(ligne)
                rows.append(values)
        feuilles.append({"name": sheet_name, "rows": rows})
    return {"text": "\n".join(lignes),"pages": [],"page_count": None,"sheets": feuilles,}

def _extraire_image(file_path: str) -> Dict[str, Any]:
    image = cv2.imread(file_path)
    if image is None:
        return {"text": "","pages": [],"page_count": 1,}
    engine = _get_onnx_ocr_engine()
    if engine is None:
        return {"text": "","pages": [],"page_count": 1}
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    texte = _ocr_image(engine, gray)
    return {
        "text": texte,
        "pages": [{"page_number": 1,"text": texte,"method": "RAPIDOCR_ONNX_FULL_IMAGE",}],
        "page_count": 1,
    }

def extraire_document_complet_pour_rag(file_path: str) -> Dict[str, Any]:
    """
    Extracteur indépendant du Learning Service. Normalise les documents DOC/DOCX en PDF via LibreOffice afin d'utiliser ensuite le même pipeline d'extraction PDF.
    """
    if not file_path:
        return {"text": "","pages": [],"detected_types": [],"source_type": None,}
    path = Path(file_path)
    if not path.exists():
        logger.error("[RAG] Fichier introuvable : %s",file_path)
        return { "text": "", "pages": [], "detected_types": [], "source_type": None, }
    extension = path.suffix.lower()
    logger.info("[RAG-EXTRACT] Début extraction complète : %s",path.name)   
    source_type = None
    try:
        # PDF
        if extension == ".pdf":
            result = _extraire_pdf(file_path)
            source_type = "PDF"
        # DOC / DOCX
        elif extension in [".doc", ".docx"]:
            logger.info("[RAG-EXTRACT] DOC/DOCX détecté : %s",path.name)
            # La conversion persistante sera faite plus tard
            # par traiter_extraction_rag(), dans :
            # rag_extracted/<tender_relative_path>/_sources/
            source_type = ("DOC" if extension == ".doc" else "DOCX")
            result = {"text": "","pages": [],"page_count": None,}
            logger.info("[RAG-EXTRACT] Conversion DOC/DOCX différée vers le dossier RAG du tender.")
        # EXCEL
        elif extension in [".xls", ".xlsx", ".xlsm"]:
            result = _extraire_excel(file_path)
            source_type = "EXCEL"
        # IMAGES
        elif extension in [".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"]:
            result = _extraire_image(file_path)
            source_type = "IMAGE"
        # NON SUPPORTÉ
        else:
            logger.warning("[RAG] Extension non supportée : %s",extension)
            result = {"text": "","pages": [],"page_count": None,}
            source_type = extension
        # POST-TRAITEMENT COMMUN

        texte = result.get("text", "")
        detected_types = detecter_types_documents(texte)
        result["detected_types"] = detected_types
        result["source_type"] = source_type
        result["source_file"] = str(file_path)

        logger.info("[RAG-EXTRACT] Extraction terminée : %s caractères | types détectés=%s",len(texte),detected_types)
        return result

    except Exception as e:
        logger.error("[RAG-EXTRACT] Erreur extraction %s : %s",file_path,e,exc_info=True)
        return {"text": "","pages": [],"detected_types": [],"source_type": source_type,"source_file": str(file_path),}

def detecter_modeles_sur_page(text: str) -> List[str]:
    """
    Détecte uniquement les modèles/formulaires présents sur une page.
    Contrairement à detecter_types_documents(), cette fonction ne considère pas RC/CPS comme des modèles.
    """
    text_norm = normaliser_texte_detection(text)
    modeles = []

    for type_document in TYPES_MODELES:
        patterns = DOCUMENT_PATTERNS.get(type_document, [])
        for pattern in patterns:
            if re.search(pattern, text_norm, re.IGNORECASE):
                modeles.append(type_document)
                break
    return modeles

def _preparer_pdf_pour_extraction_modeles(file_path: str,output_dir: Optional[str] = None,) -> Optional[str]:
    """
    Prépare une source PDF persistante pour : - extraction des modèles/formulaires - détection des zones administratives
    PDF : -> PDF original
    DOC/DOCX : -> conversion LibreOffice -> PDF sauvegardé dans rag_extracted
    IMPORTANT : aucune conversion temporaire.
    """

    path = Path(file_path)
    extension = path.suffix.lower()

    # PDF
    if extension == ".pdf":
        logger.info("[RAG-MODELES] Source déjà PDF : %s",file_path)
        return str(path)

    # DOC / DOCX
    if extension in [".doc", ".docx"]:
        if not output_dir:
            logger.error("[RAG-MODELES] output_dir obligatoire pour " "convertir un DOC/DOCX.")
            return None
        output_dir_path = Path(output_dir)
        output_dir_path.mkdir(parents=True,exist_ok=True)
        logger.info("[RAG-MODELES] Conversion persistante %s -> PDF | output=%s",extension,output_dir_path)
        converted_pdf = _convertir_avec_libreoffice(file_path=file_path,output_dir=str(output_dir_path))

        if converted_pdf and os.path.exists(converted_pdf):
            logger.info("[RAG-MODELES] PDF normalisé sauvegardé : %s",converted_pdf)
            return converted_pdf
        logger.error("[RAG-MODELES] Échec conversion %s -> PDF : %s",extension,file_path)
        return None
    logger.warning("[RAG-MODELES] Type non convertible en PDF : %s",extension)
    return None

def extraire_pages_modeles_pdf(
    source_pdf: str,
    output_dir: str,
    types_a_extraire: Optional[List[str]] = None,
    contexte_pages: int = 0,
) -> List[Dict[str, Any]]:
    """
    Extraction stricte des modèles/formulaires.
    IMPORTANT :
    - Une simple mention d'un modèle dans une page ne suffit PAS.
    - Le début doit être confirmé par page_est_vrai_modele().
    - La fin est déterminée par trouver_page_fin_modele().
    - On n'ajoute plus automatiquement +/- 1 page autour du résultat.
    - Cela évite notamment :
        * faux DECLARATION_HONNEUR
        * BDP démarrant plusieurs pages trop tôt
        * perte des pages de continuation d'un ACTE_ENGAGEMENT
    """
    os.makedirs(output_dir, exist_ok=True)
    resultats = []
    doc = fitz.open(source_pdf)
    try:
        # 1. EXTRACTION TEXTE PAGE PAR PAGE
        pages_textes: List[str] = []
        for page_index in range(len(doc)):
            page = doc.load_page(page_index)
            texte = page.get_text("text",sort=True).strip()
            extraction_method = "PYMUPDF"
            # Même logique que ton extraction principale :
            # si le texte natif est absent OU mauvais, on utilise l'OCR complet.
            if not texte or _ocr_est_de_mauvaise_qualite(texte):                
                logger.info("[RAG-MODELES] Page %d : texte natif absent/faible -> OCR.",page_index + 1)
                texte_ocr = extraire_page_complete_fast_ocr(page,page_index + 1)
                if len(texte_ocr) > len(texte):
                    texte = texte_ocr
                    extraction_method = "RAPIDOCR_ONNX_FULL_PAGE"
            pages_textes.append(texte)
            logger.debug("[RAG-MODELES] Page %d | méthode=%s | caractères=%d",page_index + 1,extraction_method,len(texte))
        # 2. TYPES AUTORISÉS
        if types_a_extraire:
            types_cibles = [t for t in types_a_extraire if t in TYPES_MODELES]
        else:
            types_cibles = list(TYPES_MODELES)
        logger.info("[RAG-MODELES] Types à rechercher strictement : %s",types_cibles)
        # 3. RECHERCHE STRICTE DU DÉBUT DE CHAQUE MODÈLE
        plages_detectees = []
        for model_type in types_cibles:
            plage = determiner_plage_modele(pages=pages_textes,model_type=model_type)
            if plage is None:
                logger.info("[RAG-MODELES] Aucun modèle réel détecté : %s",model_type)
                continue
            start_index, end_index = plage
            plages_detectees.append({"type": model_type,"start": start_index,"end": end_index})
            logger.info("[RAG-MODELES] Modèle confirmé : %s | pages %d -> %d",model_type,start_index + 1,end_index + 1)
        # 4. ÉVITER LES CHEVAUCHEMENTS
        # On trie les modèles selon leur page de début.
        plages_detectees.sort(key=lambda x: x["start"])
        plages_finales = []
        for index, plage in enumerate(plages_detectees):
            start = plage["start"]
            end = plage["end"]
            # Si un autre vrai modèle commence avant notre fin, on coupe le modèle courant juste avant celui-ci.
            if index + 1 < len(plages_detectees):
                prochain_start = plages_detectees[index + 1]["start"]
                if prochain_start <= end:
                    logger.warning(
                        "[RAG-MODELES] Chevauchement détecté : "
                        "%s pages %d-%d avec modèle suivant page %d. " "Découpage automatique.",
                        plage["type"],start + 1,end + 1,prochain_start + 1)
                    end = prochain_start - 1

            if end >= start:
                plages_finales.append({"type": plage["type"],"start": start,"end": end})
        # 5. CRÉATION DES PDF
        for plage in plages_finales:
            type_document = plage["type"]
            start_index = plage["start"]
            end_index = plage["end"]
            output_pdf = (Path(output_dir) / f"{type_document}.pdf")
            nouveau_pdf = fitz.open()
            try:
                for page_index in range(start_index,end_index + 1):
                    nouveau_pdf.insert_pdf(doc,from_page=page_index,to_page=page_index)
                # Sauvegarde
                nouveau_pdf.save(str(output_pdf))
            finally:
                nouveau_pdf.close()
            pages_extraites = list(range(start_index + 1,end_index + 2))
            item = {
                "type": type_document,
                "path": str(output_pdf),
                "pages": pages_extraites,
                "start_page": start_index + 1,
                "end_page": end_index + 1,
                "page_count": len(pages_extraites),
            }

            resultats.append(item)
            logger.info(
                "[RAG-EXTRACT] Modèle extrait : %s | " "pages=%s | fichier=%s",
                type_document,pages_extraites,output_pdf
            )
    finally:
        doc.close()
    return resultats
    
def extraire_bdp_excel_vers_pdf(file_path: str, output_pdf: str) -> Optional[str]:
    """
    Transforme toutes les feuilles Excel identifiées comme BDP
    en PDF exploitable par le workflow de génération.
    """
    extension = Path(file_path).suffix.lower()
    engine = None
    if extension == ".xls":
        engine = "xlrd"
    elif extension in [".xlsx", ".xlsm"]:
        engine = "openpyxl"
    try:
        sheets = pd.read_excel(file_path, sheet_name=None, header=None, engine=engine, dtype=object)
    except Exception as e:
        logger.error("[RAG-BDP] Impossible de lire Excel : %s",e,exc_info=True)
        return None

    styles = getSampleStyleSheet()
    normal_style = ParagraphStyle("BDPNormal",parent=styles["Normal"], fontSize=7, leading=9,)
    title_style = ParagraphStyle("BDPTitle",parent=styles["Title"], alignment=TA_CENTER, fontSize=14, leading=17,)
    doc = SimpleDocTemplate(output_pdf, pagesize=landscape(A4), rightMargin=20, leftMargin=20, topMargin=20,bottomMargin=20,)
    story = []
    story.append(Paragraph("BORDEREAU DES PRIX / DÉTAIL ESTIMATIF",title_style))
    story.append(Spacer(1, 15))
    feuille_bdp_trouvee = False

    for sheet_name, dataframe in sheets.items():
        # Convertir en texte pour détecter le BDP
        contenu_feuille = " ".join(str(v) for v in dataframe.astype(str).values.flatten())
        types = detecter_types_documents(contenu_feuille)
        est_bdp = ("BDP" in types or "BPU" in types or "DPGF" in types)
        if not est_bdp:
            continue
        feuille_bdp_trouvee = True
        story.append(Paragraph(f"Feuille : {sheet_name}",styles["Heading2"]))
        data = []

        for _, row in dataframe.iterrows():
            values = []
            for value in row.tolist():
                if pd.isna(value):
                    value = ""
                value = str(value).strip()
                values.append(Paragraph(value.replace("&","&amp;"),normal_style))
            if any(str(x.text).strip() for x in values):
                data.append(values)
        if data:
            table = Table(data,repeatRows=1)
            table.setStyle(TableStyle([
                        ("GRID",(0, 0),(-1, -1), 0.4, colors.grey,),
                        ("VALIGN",(0, 0),(-1, -1), "TOP",),
                        ("FONTNAME",(0, 0),(-1, 0),"Helvetica-Bold",),
                        ("BACKGROUND",(0, 0),(-1, 0),colors.lightgrey,),
                    ]))
            story.append(table)
            story.append(PageBreak())
    if not feuille_bdp_trouvee:
        logger.warning("[RAG-BDP] Aucun BDP détecté dans %s", file_path)
        return None
    doc.build(story)
    logger.info("[RAG-BDP] PDF généré : %s", output_pdf)
    return output_pdf

def traiter_extraction_rag(file_path: str, extracted_root: str, tender_relative_path: str) -> Dict[str, Any]:
    """
    Pipeline complet d'extraction documentaire RAG.
    Pipeline :
    1. Extraction du contenu source
    2. Détection des types
    3. Préparation d'un PDF normalisé pour DOC/DOCX
    4. Extraction des modèles/formulaires
    5. Extraction BDP Excel
    6. Détection des zones administratives
    7. Retour du texte intégral
    """
    # 1. EXTRACTION DU CONTENU SOURCE
    extraction = extraire_document_complet_pour_rag(file_path)
    texte_complet = extraction.get("text", "")
    detected_types = extraction.get("detected_types", [])
    zones_administratives = []
    extension = Path(file_path).suffix.lower()
    # 2. DOSSIER DE SORTIE
    output_base = (Path(extracted_root) / tender_relative_path)
    # Sécurité : aucun fichier ne doit être écrit directement
    # dans la racine rag_extracted.
    if output_base == Path(extracted_root).resolve():
        logger.error("[RAG-EXTRACT] tender_relative_path vide ou invalide : %r", tender_relative_path)
        raise ValueError(f"tender_relative_path invalide : {tender_relative_path!r}")

    output_base.mkdir(parents=True, exist_ok=True)
    fichiers_extraits = []
    # 3. PRÉPARATION DU PDF NORMALISÉ
    pdf_source_modeles = None
    if extension in [".pdf", ".doc", ".docx"]:
        pdf_source_modeles = (
            _preparer_pdf_pour_extraction_modeles(file_path=file_path, output_dir=str(output_base / "_sources"))
        )
        if pdf_source_modeles and extension in [".doc", ".docx"]:
            logger.info("[RAG-EXTRACT] Ré-extraction du PDF normalisé : %s", pdf_source_modeles)
            extraction_pdf = _extraire_pdf(pdf_source_modeles)
            extraction["text"] = extraction_pdf.get("text", "")
            extraction["pages"] = extraction_pdf.get("pages", [])
            extraction["page_count"] = extraction_pdf.get("page_count")
            extraction["detected_types"] = detecter_types_documents(extraction["text"])
            texte_complet = extraction["text"]
            detected_types = extraction["detected_types"]
            logger.info("[RAG-EXTRACT] PDF normalisé ré-extrait : %d caractères", len(texte_complet))
        logger.info("[RAG-MODELES] Source PDF pour analyse : %s", pdf_source_modeles)
    # 4. DÉTECTION DES ZONES ADMINISTRATIVES
    if pdf_source_modeles:
        zones_administratives = extraire_zones_administratives_pdf_complet(pdf_source_modeles)
    extraction["administrative_zones"] = zones_administratives
    # 5. EXTRACTION DES MODÈLES
    if pdf_source_modeles:
        logger.info("[RAG-MODELES] Analyse des modèles : %s", pdf_source_modeles)
        # On ne demande QUE les vrais modèles.
        types_modeles_detectes = [t for t in detected_types if t in TYPES_MODELES]
        logger.info("[RAG-MODELES] Types candidats depuis extraction globale : %s", types_modeles_detectes)
        model_files = extraire_pages_modeles_pdf(
            source_pdf=pdf_source_modeles,
            output_dir=str(output_base),
            types_a_extraire=types_modeles_detectes if types_modeles_detectes else None,
            contexte_pages=0,
        )
        # 6. PLACEMENT DANS LES DOSSIERS
        for item in model_files:
            type_document = item["type"]
            source = Path(item["path"])
            destination_dir = (output_base / type_document)
            destination_dir.mkdir(parents=True, exist_ok=True)
            destination = (destination_dir / f"{type_document}.pdf")

            if source.exists():
                if destination.exists():
                    destination.unlink()
                shutil.move(str(source),str(destination))
                item["path"] = str(destination)
                fichiers_extraits.append(item)
    # 7. BDP EXCEL
    elif extension in [".xls", ".xlsx", ".xlsm"]:
        contenu = normaliser_texte_detection(texte_complet)
        types_excel = detecter_types_documents(contenu)
        if any(x in types_excel for x in ["BDP", "BPU", "DPGF"]):
            dossier_bdp = (output_base / "BDP")
            dossier_bdp.mkdir(parents=True,exist_ok=True)
            output_pdf = (dossier_bdp / "BDP.pdf")
            pdf_bdp = extraire_bdp_excel_vers_pdf(file_path,str(output_pdf))
            if pdf_bdp:
                fichiers_extraits.append({"type": "BDP","path": pdf_bdp})
    # 8. MÉTADONNÉES FINALES
    extraction["extracted_files"] = fichiers_extraits
    extraction["extracted_root"] = str(output_base)
    extraction["administrative_zone_count"] = len(zones_administratives)
    logger.info(
        "[RAG-EXTRACT] Extraction terminée : %s | " 
        "%s fichiers modèles générés | " "%s zones administratives détectées",
        Path(file_path).name,len(fichiers_extraits),len(zones_administratives),
    )
    return extraction
   
def normaliser_texte_zone(text: str) -> str:
    """
    Normalisation légère utilisée uniquement pour la détection des zones administratives.
    """
    if not text:
        return ""
    text = text.lower()
    replacements = {"é": "e","è": "e","ê": "e","ë": "e","à": "a", "â": "a","î": "i","ï": "i","ô": "o","ù": "u","û": "u","ç": "c",}

    for src, dst in replacements.items():
        text = text.replace(src, dst)

    text = re.sub(r"\s+", " ", text)
    return text.strip()    
    
def determiner_type_zone(text: str) -> Optional[str]:
    text_norm = normaliser_texte_zone(text)

    if not text_norm:
        return None
    # 1. SIGNATURE / CACHET
    signature_patterns = [
        r"\bsignature\s+et\s+cachet\b",
        r"\bcachet\s+et\s+signature\b",
        r"\bsign[ée]\s+par\b",
        r"\bsign[ée]\s+du\b",
        r"\bsignature\b",
        r"\bcachet\b",
    ]

    for pattern in signature_patterns:
        if re.search(pattern, text_norm):
            return "SIGNATURE_CACHET"
    # 2. VISA / APPROBATION
    validation_patterns = [
        r"\bvu\s+et\s+v[ée]rifi[ée]\s+par\b",
        r"\bvu\s+et\s+pr[ée]sent[ée]\s+par\b",
        r"\blu\s+et\s+accept[ée]\s+par\b",
        r"\blu\s+et\s+approuv[ée]\s+par\b",
        r"\bvis[ée]\s+par\b",
        r"\bvisa\s+(du|de|par)\b",
        r"\bapprouv[ée]\s+par\b",
    ]

    for pattern in validation_patterns:
        if re.search(pattern, text_norm):
            return "VISA_APPROBATION"
    # 3. DATE / LIEU
    date_patterns = [
        r"\bfait\s+[àa]\b.*\ble\b",
        r"\b[a-zà-ÿ]+,\s*le\s*:",
        r"\ba\s+.{0,80},\s*le\s*:",
        r"\ble\s*:\s*\.{2,}",
    ]

    for pattern in date_patterns:
        if re.search(pattern, text_norm):
            return "DATE_LIEU"
    return None   
    
def detecter_rectangles_page(page: fitz.Page) -> List[fitz.Rect]:
    """
    Détecte les rectangles / grandes zones graphiques présents sur une page PDF.
    Utile pour identifier les cases de signature, visa, approbation, etc.
    """
    rectangles = []
    try:
        drawings = page.get_drawings()
        for drawing in drawings:
            rect = drawing.get("rect")
            if not rect:
                continue
            rect = fitz.Rect(rect)
            # Éliminer les petits éléments graphiques
            if rect.width < 80 or rect.height < 30:
                continue
            # Éliminer les énormes rectangles correspondant éventuellement à toute la page
            if rect.width > page.rect.width * 0.98 and rect.height > page.rect.height * 0.98:
                continue
            rectangles.append(rect)
    except Exception as e:
        logger.warning("[RAG][ZONE] Impossible de détecter les rectangles : %s", e)
    return rectangles

def detecter_zones_administratives_pdf(page: fitz.Page,page_number: int) -> List[Dict[str, Any]]:
    """
    Détecte les zones administratives importantes d'une page :
    - signature- cachet- visa- approbation- lu et accepté- fait à / le- zones de validation
    Retourne également la page et les coordonnées.
    """
    zones = []

    # 1. Récupération des blocs texte
    text_dict = page.get_text("dict")
    text_blocks = []
    for block in text_dict.get("blocks", []):
        if block.get("type") != 0:
            continue
        block_text_parts = []
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                txt = span.get("text", "").strip()
                if txt:
                    block_text_parts.append(txt)
        block_text = " ".join(block_text_parts).strip()
        if not block_text:
            continue
        bbox = fitz.Rect(block["bbox"])
        text_blocks.append({"text": block_text,"bbox": bbox})

    # 2. Détection des rectangles
    rectangles = detecter_rectangles_page(page)
    # 3. Recherche des libellés administratifs
    for block in text_blocks:
        zone_type = determiner_type_zone(block["text"])
        if not zone_type:
            continue
        label_rect = block["bbox"]
        # Recherche d'une case contenant le libellé
        meilleure_case = None
        meilleur_score = 0
        for rect in rectangles:
            # Le rectangle doit contenir ou être proche du libellé.
            centre_x = rect.x0 + rect.width / 2
            centre_y = rect.y0 + rect.height / 2
            if rect.contains(label_rect):
                score = 100
            else:
                distance = abs(centre_y - (label_rect.y0 + label_rect.height / 2))
                if distance < 150:
                    score = max(0, 80 - distance / 2)
                else:
                    score = 0
            if score > meilleur_score:
                meilleur_score = score
                meilleure_case = rect

        # Si une case a été trouvée
        if meilleure_case:
            zone_rect = meilleure_case
        else:
            # Fallback : zone autour du texte
            zone_rect = fitz.Rect(max(0, label_rect.x0 - 20),max(0, label_rect.y0 - 20),min(page.rect.width, label_rect.x1 + 200),min(page.rect.height, label_rect.y1 + 150))
        zones.append({
            "page_number": page_number,
            "type": zone_type,
            "label": block["text"],
            "bbox": [round(zone_rect.x0, 2),round(zone_rect.y0, 2),round(zone_rect.x1, 2),round(zone_rect.y1, 2)],
            "source": "PYMUPDF_TEXT_AND_DRAWINGS",
            "is_signature_zone": zone_type == "SIGNATURE_CACHET",
            "is_validation_zone": zone_type == "VISA_APPROBATION",
            "is_date_zone": zone_type == "DATE_LIEU",
        })
    return zones

def ocr_page_complete_avec_positions(page: fitz.Page, zoom: float = 2.0) -> List[Dict[str, Any]]:
    """
    OCR complet d'une page avec conservation des coordonnées des textes détectés.
    Utilisé uniquement pour les documents scannés.
    """
    resultats = []
    try:
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat,alpha=False)
        img = np.frombuffer(pix.samples,dtype=np.uint8).reshape(pix.height,pix.width,3)
        gray = cv2.cvtColor(img,cv2.COLOR_BGR2GRAY)
        engine = _get_onnx_ocr_engine()
        if engine is None:
            return []
        result, _ = engine(gray)
        if not result:
            return []
        for item in result:
            if len(item) < 3:
                continue
            points = item[0]
            text = str(item[1]).strip()
            confidence = float(item[2])
            if not text:
                continue
            if confidence < 0.40:
                continue
            xs = [float(point[0]) for point in points]
            ys = [float(point[1]) for point in points]
            x0 = min(xs) / zoom
            y0 = min(ys) / zoom
            x1 = max(xs) / zoom
            y1 = max(ys) / zoom

            resultats.append({
                "text": text,
                "confidence": confidence,
                "bbox": [round(x0, 2),round(y0, 2),round(x1, 2),round(y1, 2)]
            })
    except Exception as e:
        logger.error("[RAG][OCR] Erreur OCR page complète : %s",e,exc_info=True)
    return resultats    
    
def detecter_zones_administratives_scan(page: fitz.Page,page_number: int) -> List[Dict[str, Any]]:
    ocr_items = ocr_page_complete_avec_positions(page)

    if not ocr_items:
        return []
    zones = []
    for item in ocr_items:
        text = item["text"]
        zone_type = determiner_type_zone(text)
        if not zone_type:
            continue
        bbox = item["bbox"]

        # Agrandir la zone autour du libellé pour couvrir la case de signature/visa.
        x0, y0, x1, y1 = bbox
        zone_rect = fitz.Rect(max(0, x0 - 30),max(0, y0 - 30),min(page.rect.width, x1 + 250),min(page.rect.height, y1 + 180))
        zones.append({
            "page_number": page_number,
            "type": zone_type,
            "label": text,
            "bbox": [round(zone_rect.x0, 2),round(zone_rect.y0, 2),round(zone_rect.x1, 2),round(zone_rect.y1, 2)],
            "source": "RAPIDOCR_ONNX",
            "confidence": item["confidence"],
            "is_signature_zone": zone_type == "SIGNATURE_CACHET",
            "is_validation_zone": zone_type == "VISA_APPROBATION",
            "is_date_zone": zone_type == "DATE_LIEU",
        })
    return zones
    
def detecter_zones_administratives(page: fitz.Page,page_number: int) -> List[Dict[str, Any]]:
    # Tentative 1 : texte natif
    zones = detecter_zones_administratives_pdf(page=page,page_number=page_number)
    if zones:
        logger.info("[RAG][ZONES] Page %s : %s zone(s) détectée(s) via texte natif.",page_number,len(zones))
        return zones
    # Tentative 2 : document scanné
    zones = detecter_zones_administratives_scan(page=page,page_number=page_number)
    if zones:
        logger.info("[RAG][ZONES] Page %s : %s zone(s) détectée(s) via OCR.",page_number,len(zones))
    return zones
    
def extraire_zones_administratives_pdf_complet(file_path: str) -> List[Dict[str, Any]]:
    """
    Analyse toutes les pages d'un PDF afin de détecter les zones administratives utiles au workflow de génération : signatures - cachets - visas - approbations - acceptations - dates / lieux - zones de validation
    Pour chaque zone, on conserve : page_number - type - label - bbox - source - confiance éventuelle
    """

    zones_detectees = []
    try:
        doc = fitz.open(file_path)
        logger.info("[RAG][ZONES] Analyse administrative : %s pages | %s", len(doc), Path(file_path).name)

        for page_index in range(len(doc)):
            page = doc[page_index]
            page_number = page_index + 1
            zones_page = detecter_zones_administratives(page=page, page_number=page_number)

            if zones_page:
                zones_detectees.extend(zones_page)
                logger.info("[RAG][ZONES] Page %s : %s zone(s)", page_number, len(zones_page))

        doc.close()

    except Exception as e:
        logger.error("[RAG][ZONES] Erreur analyse zones administratives : %s", e, exc_info=True)

    logger.info("[RAG][ZONES] Total détecté : %s zone(s)", len(zones_detectees))
    return zones_detectees